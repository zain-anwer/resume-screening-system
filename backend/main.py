"""
FastAPI backend entrypoint for the Resume Screening & Ranking system.

Orchestrates the four pipeline stages already implemented under `src/`:

    1. ingestion          -> src.ingestion.resume_ingestion.ingest_resumes
    2. field extraction   -> src.extraction.field_extraction.extract_fields
    3. eligibility check  -> src.policy_engine.candidate_evaluation.evaluate_candidates
    4. ranking            -> src.retrieval.resume_ranking.rank_candidates

and reshapes their JSON outputs into the exact response shapes the
frontend's src/api/client.js expects (see data/mockData.js), plus a
merged "full candidate record" endpoint (see /api/candidates/process
below) that combines all four stages into one object per candidate.

Run with:
    uvicorn main:app --reload --port 8000

KNOWN GAPS IN THE CURRENT REPO (see chat for details):
  - `models/` package (Candidate, JobDescription, RankedCandidate dataclasses)
    is imported by several stage-4 modules but wasn't present in the
    uploaded src/. Stages 1-3 run fine without it; ranking will raise
    ImportError until it's added. Handled below via a lazy/guarded import
    so the rest of the API still works.
  - `configs/<job_category>.yaml` rule files (used by candidate_evaluation.py)
    are expected at <project_root>/configs/ and weren't in the upload.
"""

from __future__ import annotations

import json
import uuid
from pathlib import Path
from pydantic import BaseModel
from typing import Optional

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from docx import Document
from src.ingestion.resume_ingestion import ingest_resumes
from src.extraction.field_extraction import extract_fields
from src.policy_engine.candidate_evaluation import evaluate_candidates

# Ranking depends on the currently-missing `models` package. Import it
# lazily so a broken stage 4 doesn't take down ingestion/extraction/
# eligibility, which don't need it.
try:
    from src.retrieval.resume_ranking import rank_candidates

    RANKING_AVAILABLE = True
    _ranking_import_error: Optional[str] = None
except ImportError as exc:  # pragma: no cover - depends on repo state
    rank_candidates = None
    RANKING_AVAILABLE = False
    _ranking_import_error = str(exc)


# ---------------------------------------------------------------------
# App setup
# ---------------------------------------------------------------------

RUNS_DIR = Path("runs")
RUNS_DIR.mkdir(exist_ok=True)

app = FastAPI(title="Resume Screening API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # tighten to your frontend origin before shipping
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory pointer to the most recently completed run. Swap for a real
# store (DB / redis / per-run files keyed by run_id) once you have more
# than one user or want history across server restarts.
_latest_run: dict = {}


def _run_dir(run_id: str) -> Path:
    d = RUNS_DIR / run_id
    d.mkdir(parents=True, exist_ok=True)
    return d


def _require_run() -> dict:
    if not _latest_run:
        raise HTTPException(
            status_code=404,
            detail="No pipeline run yet. POST /api/pipeline/run or /api/candidates/process first.",
        )
    return _latest_run


def _policy_label(overall_status: str) -> str:
    return {"Eligible": "Pass", "Not Eligible": "Fail"}.get(overall_status, "Review")


def _inject_eligibility_status(extracted: list, eligibility: list) -> list[dict]:
    """
    Stamps each extracted candidate record with an "overall_status" key
    (e.g. "Eligible" / "Not Eligible") pulled from the eligibility stage,
    keyed on candidate_id <-> id. This is the file the ranking module now
    reads, so it can filter out "Not Eligible" candidates before scoring.

    Returns a new list — the original `extracted` objects aren't mutated,
    since that list is also used later to build the merged response.
    """
    status_by_id = {e["candidate_id"]: e["overall_status"] for e in eligibility}

    stamped = []
    for candidate in extracted:
        record = dict(candidate)
        record["overall_status"] = status_by_id.get(candidate["id"])
        stamped.append(record)
    return stamped


def _run_full_pipeline(
    folder_path: str, job_description_path: str, top_k: int
) -> tuple[str, list, list, list, Optional[str]]:
    """
    Runs all four stages back to back and returns:
        (run_id, extracted, eligibility, ranked, ranking_error)

    Raises HTTPException on bad input or a missing/unusable job description,
    since the merged endpoint requires ranking (unlike /api/pipeline/run,
    where ranking is optional).
    """
    folder = Path(folder_path)
    if not folder.exists() or not folder.is_dir():
        raise HTTPException(status_code=400, detail=f"Folder not found: {folder_path}")

    jd_path = Path(job_description_path)
    if not jd_path.exists():
        raise HTTPException(
            status_code=400, detail=f"Job description not found: {job_description_path}"
        )

    if not RANKING_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail=f"Ranking stage unavailable: {_ranking_import_error}",
        )

    run_id = uuid.uuid4().hex[:12]
    out = _run_dir(run_id)

    ingested_path = out / "01_ingested.json"
    extracted_path = out / "02_extracted.json"
    eligibility_path = out / "03_eligibility.json"
    extracted_for_ranking_path = out / "03b_extracted_with_status.json"
    ranked_path = out / "04_ranked.json"

    # Stage 1: ingestion (OCR/text extraction from raw resume + CNIC files)
    ingest_resumes(str(folder), str(ingested_path))

    # Stage 2: field extraction (NER/regex parsing into structured profile JSON)
    extract_fields(str(ingested_path), str(extracted_path))

    # Stage 3: eligibility / policy evaluation against YAML rules
    evaluate_candidates(str(extracted_path), str(eligibility_path))

    extracted = json.loads(extracted_path.read_text(encoding="utf-8"))
    eligibility = json.loads(eligibility_path.read_text(encoding="utf-8"))

    # Stamp each candidate with "overall_status" from the eligibility stage
    # and write that as the ranking module's input — it now filters out
    # "Not Eligible" candidates before scoring.
    extracted_with_status = _inject_eligibility_status(extracted, eligibility)
    extracted_for_ranking_path.write_text(
        json.dumps(extracted_with_status, indent=2), encoding="utf-8"
    )

    # Stage 4: ranking against the supplied job description
    ranking_error: Optional[str] = None
    ranked: list = []
    try:
        rank_candidates(
            str(jd_path),
            str(extracted_for_ranking_path),
            str(ranked_path),
            top_k=top_k,
        )
        ranked = json.loads(ranked_path.read_text(encoding="utf-8"))
    except Exception as exc:  # noqa: BLE001 - surface any ranking failure, don't crash the run
        ranking_error = f"Ranking failed: {exc}"

    return run_id, extracted, eligibility, ranked, ranking_error


def _merge_candidate_records(extracted: list, eligibility: list, ranked: list) -> list[dict]:
    """
    Combines the three per-candidate JSON outputs (extraction, eligibility,
    ranking) into a single object per candidate, keyed on candidate_id.

    - `extracted` supplies the full profile (personal_info, education,
      experience, skills, certifications, etc.) — this is the base object.
    - `eligibility` is nested under "eligibility" (None if a candidate
      wasn't evaluated, which shouldn't normally happen).
    - `ranking` is nested under "ranking" (None for candidates outside
      top_k, since the ranking stage only returns its top results).
    """
    eligibility_by_id = {e["candidate_id"]: e for e in eligibility}
    ranking_by_id = {r["candidate_id"]: r for r in ranked}

    merged: list[dict] = []
    for candidate in extracted:
        cid = candidate["id"]
        elig = eligibility_by_id.get(cid)
        rank = ranking_by_id.get(cid)

        record = dict(candidate)  # shallow copy of the full extracted profile
        record["eligibility"] = (
            {
                "overall_status": elig["overall_status"],
                "policy_label": _policy_label(elig["overall_status"]),
                "education": elig["education"],
                "experience": elig["experience"],
                "supervisory": elig["supervisory"],
                "age": elig["age"],
                "preferred_skills": elig["preferred_skills"],
            }
            if elig
            else None
        )
        record["ranking"] = (
            {
                "rank": rank["rank"],
                "lexical_score": rank["lexical_score"],
                "semantic_score": rank["semantic_score"],
                "final_score": rank["final_score"],
                "matched_skills": rank["matched_skills"],
                "missing_skills": rank["missing_skills"],
                "matched_preferred_skills": rank["matched_preferred_skills"],
            }
            if rank
            else None
        )
        merged.append(record)

    # Candidates who made the ranked top_k float to the top, ordered by rank;
    # everyone else follows in their original extraction order.
    merged.sort(key=lambda r: r["ranking"]["rank"] if r["ranking"] else float("inf"))
    return merged


# ---------------------------------------------------------------------
# Request/response models
# ---------------------------------------------------------------------

class ProcessRequest(BaseModel):
    folder_path: str = Field(
        ...,
        description=(
            "Path to the jobs folder, structured as "
            "<folder_path>/<job_category>/<candidate_folder>/{resume, cnic}"
        ),
    )
    job_description_path: Optional[str] = Field(
        None,
        description="Path to a job description .txt file. Ranking is skipped if omitted.",
    )
    top_k: int = 20


class ProcessResponse(BaseModel):
    run_id: str
    candidates_processed: int
    ranking_available: bool
    ranking_error: Optional[str] = None


class MergeRequest(BaseModel):
    folder_path: str = Field(
        ...,
        description=(
            "Path to the jobs folder, structured as "
            "<folder_path>/<job_category>/<candidate_folder>/{resume, cnic}"
        ),
    )
    job_description_path: str = Field(
        ...,
        description="Path to a job description .txt file. Required — ranking runs as part of this call.",
    )
    top_k: int = 20


# ---------------------------------------------------------------------
# Pipeline run (existing behaviour — ranking optional)
# ---------------------------------------------------------------------

@app.post("/api/pipeline/run", response_model=ProcessResponse)
def run_pipeline(req: ProcessRequest) -> ProcessResponse:
    folder = Path(req.folder_path)
    if not folder.exists() or not folder.is_dir():
        raise HTTPException(status_code=400, detail=f"Folder not found: {req.folder_path}")

    run_id = uuid.uuid4().hex[:12]
    out = _run_dir(run_id)

    ingested_path = out / "01_ingested.json"
    extracted_path = out / "02_extracted.json"
    eligibility_path = out / "03_eligibility.json"
    ranked_path = out / "04_ranked.json"

    ingest_resumes(str(folder), str(ingested_path))
    extract_fields(str(ingested_path), str(extracted_path))
    evaluate_candidates(str(extracted_path), str(eligibility_path))

    extracted = json.loads(extracted_path.read_text(encoding="utf-8"))
    eligibility = json.loads(eligibility_path.read_text(encoding="utf-8"))

    ranked: list = []
    ranking_error = _ranking_import_error if not RANKING_AVAILABLE else None

    if req.job_description_path:
        if not RANKING_AVAILABLE:
            ranking_error = f"Ranking stage unavailable: {_ranking_import_error}"
        elif not Path(req.job_description_path).exists():
            ranking_error = f"Job description not found: {req.job_description_path}"
        else:
            # Stamp with "overall_status" before ranking — the ranking
            # module filters out "Not Eligible" candidates on this field.
            extracted_with_status = _inject_eligibility_status(extracted, eligibility)
            extracted_for_ranking_path = out / "02b_extracted_with_status.json"
            extracted_for_ranking_path.write_text(
                json.dumps(extracted_with_status, indent=2), encoding="utf-8"
            )
            rank_candidates(
                req.job_description_path,
                str(extracted_for_ranking_path),
                str(ranked_path),
                top_k=req.top_k,
            )
            ranked = json.loads(ranked_path.read_text(encoding="utf-8"))

    global _latest_run
    _latest_run = {
        "run_id": run_id,
        "extracted": extracted,
        "eligibility": eligibility,
        "ranked": ranked,
        "ranking_error": ranking_error,
    }

    return ProcessResponse(
        run_id=run_id,
        candidates_processed=len(extracted),
        ranking_available=RANKING_AVAILABLE,
        ranking_error=ranking_error,
    )


# ---------------------------------------------------------------------
# NEW: merged full-candidate-record endpoint
# ---------------------------------------------------------------------

@app.post("/api/candidates/process")
def process_and_merge(req: MergeRequest) -> list[dict]:
    """
    Takes a folder path and a job description path, runs the full
    pipeline (ingestion -> extraction -> eligibility -> ranking), and
    returns ONE list of merged candidate objects — each object contains
    the full extracted profile plus nested "eligibility" and "ranking"
    sections. This also updates the in-memory `_latest_run` so the
    existing dashboard/ranking/queue endpoints keep working afterward.
    """
    run_id, extracted, eligibility, ranked, ranking_error = _run_full_pipeline(
        req.folder_path, req.job_description_path, req.top_k
    )

    global _latest_run
    _latest_run = {
        "run_id": run_id,
        "extracted": extracted,
        "eligibility": eligibility,
        "ranked": ranked,
        "ranking_error": ranking_error,
    }

    if ranking_error:
        raise HTTPException(status_code=502, detail=ranking_error)

    merged_path = _run_dir(run_id) / "05_merged.json"
    merged = _merge_candidate_records(extracted, eligibility, ranked)
    merged_path.write_text(json.dumps(merged, indent=2), encoding="utf-8")

    return merged


# ---------------------------------------------------------------------
# GET endpoints matching src/api/client.js
# ---------------------------------------------------------------------

@app.get("/api/dashboard/summary")
def dashboard_summary() -> dict:
    """Matches fetchDashboardData() -> dashboardMock shape."""
    run = _require_run()
    extracted = run["extracted"]
    eligibility_by_id = {e["candidate_id"]: e for e in run["eligibility"]}
    ranked_by_id = {r["candidate_id"]: r for r in run["ranked"]}

    total = len(extracted)
    screened = len(eligibility_by_id)
    shortlisted = sum(
        1 for e in eligibility_by_id.values() if e["overall_status"] == "Eligible"
    )

    scores = [r["final_score"] * 100 for r in run["ranked"]]
    avg_match = round(sum(scores) / len(scores), 1) if scores else 0.0

    kpis = [
        {"label": "Total Resumes", "value": f"{total:,}", "delta": None, "up": True},
        {"label": "Screened", "value": f"{screened:,}", "delta": None, "up": True},
        {"label": "Shortlisted", "value": f"{shortlisted:,}", "delta": None, "up": True},
        {"label": "Avg Match Score", "value": f"{avg_match}%", "delta": None, "up": True},
    ]

    pipeline = [
        {"stage": "Applied", "value": total, "dropoff": None},
        {"stage": "Screened", "value": screened, "dropoff": None},
        {"stage": "Shortlisted", "value": shortlisted, "dropoff": None},
    ]

    education_counts: dict[str, int] = {}
    for c in extracted:
        for edu in c.get("education") or []:
            level = (edu.get("degree_level") or edu.get("degree_raw") or "Unknown").strip() or "Unknown"
            education_counts[level] = education_counts.get(level, 0) + 1
    education = [
        {"name": k, "value": v}
        for k, v in sorted(education_counts.items(), key=lambda kv: -kv[1])
    ]

    skill_counts: dict[str, int] = {}
    for c in extracted:
        for skill in c.get("skills") or []:
            skill_counts[skill] = skill_counts.get(skill, 0) + 1
    skills = [
        {"skill": k, "count": v}
        for k, v in sorted(skill_counts.items(), key=lambda kv: -kv[1])[:8]
    ]

    def exp_bucket(years: float) -> str:
        if years <= 2:
            return "0-2 yrs"
        if years <= 5:
            return "2-5 yrs"
        if years <= 8:
            return "5-8 yrs"
        return "8+ yrs"

    exp_counts = {"0-2 yrs": 0, "2-5 yrs": 0, "5-8 yrs": 0, "8+ yrs": 0}
    for c in extracted:
        years = (c.get("experience_summary") or {}).get("total_experience_years", 0) or 0
        exp_counts[exp_bucket(years)] += 1
    experience_dist = [{"range": k, "count": v} for k, v in exp_counts.items()]

    recent_screenings = []
    for c in extracted[:20]:
        cid = c["id"]
        elig = eligibility_by_id.get(cid)
        rank = ranked_by_id.get(cid)
        personal = c.get("personal_info") or {}
        years = (c.get("experience_summary") or {}).get("total_experience_years", 0) or 0
        recent_screenings.append({
            "id": cid,
            "name": personal.get("name") or "Unknown",
            "email": personal.get("email") or "",
            "role": (c.get("metadata") or {}).get("job_category", ""),
            "match": round(rank["final_score"] * 100) if rank else None,
            "policy": _policy_label(elig["overall_status"]) if elig else "Review",
            "exp": f"{years}y",
            "skills": (c.get("skills") or [])[:2],
            "status": "Shortlisted" if elig and elig["overall_status"] == "Eligible" else "Screened",
        })

    return {
        "kpis": kpis,
        "pipeline": pipeline,
        "education": education,
        "skills": skills,
        "experienceDist": experience_dist,
        "recentScreenings": recent_screenings,
    }


@app.get("/api/ranking/latest")
def ranking_latest() -> list:
    """Matches fetchRankedCandidates() -> candidateRanking shape."""
    run = _require_run()
    if run["ranking_error"]:
        raise HTTPException(status_code=409, detail=run["ranking_error"])

    eligibility_by_id = {e["candidate_id"]: e for e in run["eligibility"]}
    extracted_by_id = {c["id"]: c for c in run["extracted"]}

    results = []
    for r in run["ranked"]:
        cid = r["candidate_id"]
        elig = eligibility_by_id.get(cid)
        candidate = extracted_by_id.get(cid, {})
        education = (candidate.get("education") or [{}])[0]
        years = (candidate.get("experience_summary") or {}).get("total_experience_years", 0) or 0

        results.append({
            "rank": r["rank"],
            "candidate": r["candidate_name"],
            "overall": round(r["final_score"] * 100),
            "semantic": round(r["semantic_score"] * 100),
            "bm25": round(r["lexical_score"] * 100),
            "experience": f"{years} yrs",
            "education": education.get("degree_level") or education.get("degree_raw") or "",
            "policy": _policy_label(elig["overall_status"]) if elig else "Review",
            "status": "Shortlisted" if elig and elig["overall_status"] == "Eligible" else "Screened",
        })
    return results


@app.get("/api/screening/queue")
def screening_queue() -> list:
    """Matches fetchScreeningQueue() -> screeningQueue shape."""
    run = _require_run()

    queue = []
    for c in run["extracted"]:
        flags = c.get("flags") or {}
        status = "Completed"
        if flags.get("missing_fields"):
            status = "Parsing"
        if flags.get("ner_review_needed"):
            status = "Review Needed"

        queue.append({
            "id": c["id"],
            "file": (c.get("metadata") or {}).get("candidate_id", c["id"]),
            "status": status,
        })
    return queue


@app.get("/api/health")
def health() -> dict:
    return {"status": "ok", "ranking_available": RANKING_AVAILABLE}

class PolicySaveRequest(BaseModel):
    job_name: str
    yaml: str

CONFIGS_DIR = Path(__file__).parent / "configs"

@app.post("/api/policy/save")
def save_policy(body: PolicySaveRequest):
    CONFIGS_DIR.mkdir(parents=True, exist_ok=True)

    safe_name = body.job_name.strip().lower().replace(" ", "_")
    if not safe_name:
        safe_name = "policy"

    file_path = CONFIGS_DIR / f"{safe_name}.yaml"
    file_path.write_text(body.yaml, encoding="utf-8")

    return {"path": str(file_path)}

    

ADS_DIR = Path(__file__).parent / "ads"


def _slugify(text: str) -> str:
    safe = "".join(c if c.isalnum() else "_" for c in text.strip().lower())
    while "__" in safe:
        safe = safe.replace("__", "_")
    return safe.strip("_") or "job"


class JobDescriptionRequest(BaseModel):
    job_title: str
    department: str = ""
    location: str = ""
    employment_type: str = ""
    experience_required: str = ""
    education_required: str = ""
    required_skills: list[str] = []
    preferred_skills: list[str] = []
    responsibilities: list[str] = []


@app.post("/api/jobs/save")
def save_job_description(body: JobDescriptionRequest):
    if not body.job_title.strip():
        raise HTTPException(status_code=400, detail="job_title is required")

    ADS_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()

    def add_field(label: str, value: str):
        doc.add_paragraph(f"{label}: {value}")
        doc.add_paragraph("")

    def add_list_section(label: str, items: list[str]):
        doc.add_paragraph(f"{label}:")
        for item in items:
            doc.add_paragraph(item)
        doc.add_paragraph("")

    add_field("Job Title", body.job_title)
    add_field("Department", body.department)
    add_field("Location", body.location)
    add_field("Employment Type", body.employment_type)
    add_field("Experience Required", body.experience_required)
    add_field("Education Required", body.education_required)

    if body.required_skills:
        add_list_section("Required Skills", body.required_skills)
    if body.preferred_skills:
        add_list_section("Preferred Skills", body.preferred_skills)
    if body.responsibilities:
        add_list_section("Responsibilities", body.responsibilities)

    file_path = ADS_DIR / f"{_slugify(body.job_title)}.docx"
    doc.save(str(file_path))

    return {"path": str(file_path)}